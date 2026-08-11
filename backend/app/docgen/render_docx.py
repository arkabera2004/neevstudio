"""Word renderers for the two deliverable families.

Built to reproduce the reference deliverables in `PoC Data/Our Attempt/` — the
exact page setup, table geometry, shading and fonts were measured from those
files and are hard-coded here as the fidelity target. python-docx has no public
API for cell/run shading, explicit table borders, fixed column widths or PAGE
fields, so those are emitted as raw OOXML through the helpers below.

Palette (measured):
  Family A compliance navy   #1F3864   section headings #16233F
  Family B requirement navy  #1F4E79   subtitle grey    #585858
  metadata label shading     #D9EAF7   zebra stripe     #F2F2F2
  risk High #C00000  Medium #ED7D31  Low #548235  (white bold text)
  provenance italic          #595959
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

from .models import ComplianceResult, GeneratedDoc, MatrixSection

# ── colours ──────────────────────────────────────────────────────────────────
NAVY_COMPLIANCE = RGBColor(0x1F, 0x38, 0x64)
NAVY_SECTION = RGBColor(0x16, 0x23, 0x3F)
NAVY_REQSET = RGBColor(0x1F, 0x4E, 0x79)
GREY_SUBTITLE = RGBColor(0x58, 0x58, 0x58)
GREY_ITALIC_SUB = RGBColor(0x40, 0x40, 0x40)
GREY_PROVENANCE = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY_PLACEHOLDER = RGBColor(0x88, 0x88, 0x88)

# Page geometry in EMU, taken verbatim from the reference files (Inches() would
# round and drift a fraction of a point).
MARGIN_LR_A, MARGIN_TB_A = 444500, 393700

FILL_HEADER_A = "1F3864"
FILL_HEADER_B = "1F4E79"
FILL_META_LABEL = "D9EAF7"
FILL_ZEBRA = "F2F2F2"
FILL_WHITE = "FFFFFF"

RISK_FILL = {"High": "C00000", "Medium": "ED7D31", "Low": "548235"}


# ── low-level OOXML helpers ──────────────────────────────────────────────────
def _shade_cell(cell, fill_hex: str) -> None:
    """Set a table-cell background fill (w:tcPr/w:shd)."""
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _shade_run(run, fill_hex: str) -> None:
    """Set a run-level background fill (w:rPr/w:shd) — used for legend chips."""
    rpr = run._element.get_or_add_rPr()
    shd = rpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        rpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _single_borders(table, sz: int = 4) -> None:
    """Apply single-line borders (all six edges) to a styleless table."""
    tblpr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblpr.append(borders)


def _fix_widths(table, widths_dxa: list[int]) -> None:
    """Pin column widths so Word does not recompute the layout.

    Sets the table grid, every cell's tcW, tblW (sum) and a fixed layout type —
    all four are required; omitting any one lets Word autofit and the measured
    geometry drifts.
    """
    total = sum(widths_dxa)
    tblpr = table._tbl.tblPr

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblpr.append(layout)

    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for gc, w in zip(grid.findall(qn("w:gridCol")), widths_dxa):
        gc.set(qn("w:w"), str(w))

    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_dxa):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(w))
            tcw.set(qn("w:type"), "dxa")


def _vcenter(cell) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center")
    tcpr.append(va)


def _field(paragraph, instr: str) -> None:
    """Append a Word field (e.g. PAGE, NUMPAGES) to a paragraph.

    Emits the begin / instrText / separate / end run sequence Word expects;
    the order matters — instrText must sit between begin and separate.
    """

    def _fld(kind: str):
        run = paragraph.add_run()
        run.font.size = Pt(8)
        el = OxmlElement("w:fldChar")
        el.set(qn("w:fldCharType"), kind)
        run._element.append(el)

    _fld("begin")
    run = paragraph.add_run()
    run.font.size = Pt(8)
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = f" {instr} "
    run._element.append(instr_el)
    _fld("separate")
    _fld("end")


def _style_cell_text(
    cell, text: str, *, bold=False, size=8.0, color=None, name="Calibri"
) -> None:
    """Replace a cell's text with a single styled run."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text or "")
    run.font.bold = bold
    run.font.size = Pt(size)
    if name:
        run.font.name = name
    if color is not None:
        run.font.color.rgb = color


def _new_table(doc, rows: int, cols: int):
    """A styleless table (no inherited 'Table Grid') for explicit border control."""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = None
    return table


# ── Family A — compliance & traceability matrix ──────────────────────────────
# column widths measured from the reference (dxa); sum 14440 == 10.03 in
COMPLIANCE_WIDTHS = [850, 2400, 2600, 2100, 2600, 2600, 1290]
COMPLIANCE_HEADERS = [
    "Req ID",
    "Requirement",
    "Rationale / Description",
    "Applicable Standard(s)",
    "Compliance Approach",
    "Risk / Hazard Addressed",
    "Risk Level",
]


def _landscape(doc) -> None:
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    sec.left_margin = sec.right_margin = Emu(MARGIN_LR_A)
    sec.top_margin = sec.bottom_margin = Emu(MARGIN_TB_A)


def _compliance_header_footer(doc, title: str) -> None:
    sec = doc.sections[0]
    hp = sec.header.paragraphs[0]
    hp.text = ""
    r = hp.add_run(title)
    r.font.size = Pt(8)
    r.font.color.rgb = GREY_PROVENANCE
    fp = sec.footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead = fp.add_run("Page ")
    lead.font.size = Pt(8)
    _field(fp, "PAGE")
    mid = fp.add_run(" of ")
    mid.font.size = Pt(8)
    _field(fp, "NUMPAGES")


def _legend_paragraph(doc) -> None:
    p = doc.add_paragraph()
    label = p.add_run("Risk Level Legend:  ")
    label.font.bold = True
    label.font.size = Pt(9)
    label.font.name = "Calibri"
    for i, level in enumerate(("High", "Medium", "Low")):
        if i:
            gap = p.add_run("   ")
            gap.font.size = Pt(9)
        chip = p.add_run(f"  {level}  ")
        chip.font.bold = True
        chip.font.size = Pt(9)
        chip.font.name = "Calibri"
        chip.font.color.rgb = WHITE
        _shade_run(chip, RISK_FILL[level])


def _compliance_section(doc, section: MatrixSection) -> None:
    heading = doc.add_heading(section.title, level=2)
    for r in heading.runs:
        r.font.bold = True
        r.font.size = Pt(12)
        r.font.name = "Calibri"
        r.font.color.rgb = NAVY_SECTION

    table = _new_table(doc, rows=1, cols=7)
    _single_borders(table)
    for cell, text in zip(table.rows[0].cells, COMPLIANCE_HEADERS):
        _shade_cell(cell, FILL_HEADER_A)
        _vcenter(cell)
        _style_cell_text(cell, text, bold=True, size=8.5, color=WHITE)

    for i, row in enumerate(section.rows):
        cells = table.add_row().cells
        stripe = FILL_ZEBRA if i % 2 else FILL_WHITE
        risk = row.risk_level
        values = [
            row.req_id,
            row.requirement,
            row.rationale or "—",
            row.standards or "—",
            row.compliance_approach or "—",
            row.risk_hazard or "—",
            risk or "—",
        ]
        for j, (cell, value) in enumerate(zip(cells, values)):
            if j == 6 and risk in RISK_FILL:
                _shade_cell(cell, RISK_FILL[risk])
                _style_cell_text(cell, value, bold=True, size=8.5, color=WHITE)
            else:
                _shade_cell(cell, stripe)
                is_id = j == 0
                color = GREY_PLACEHOLDER if (not row.enriched and j >= 2) else BLACK
                _style_cell_text(cell, value, bold=is_id, size=8.0, color=color)

    _fix_widths(table, COMPLIANCE_WIDTHS)
    doc.add_paragraph()  # spacer, matching the reference


def render_compliance_docx(result: ComplianceResult) -> bytes:
    """Render a compliance & traceability matrix (.docx) → bytes."""
    doc = Document()
    _landscape(doc)
    _compliance_header_footer(doc, result.doc_title)

    title = doc.add_paragraph()
    tr = title.add_run(result.doc_title)
    tr.font.bold = True
    tr.font.size = Pt(20)
    tr.font.name = "Calibri"
    tr.font.color.rgb = NAVY_COMPLIANCE

    if result.subtitle:
        sub = doc.add_paragraph()
        sr = sub.add_run(result.subtitle)
        sr.font.italic = True
        sr.font.size = Pt(11)
        sr.font.name = "Calibri"
        sr.font.color.rgb = GREY_ITALIC_SUB

    _legend_paragraph(doc)

    prov = doc.add_paragraph()
    pr = prov.add_run(
        f"Base requirements: {result.source_name}. Rationale, Applicable "
        "Standard(s), Compliance Approach, Risk / Hazard Addressed, and Risk "
        "Level added for design-control and DHF traceability."
    )
    pr.font.italic = True
    pr.font.size = Pt(8)
    pr.font.name = "Calibri"
    pr.font.color.rgb = GREY_PROVENANCE

    for section in result.sections:
        _compliance_section(doc, section)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Family B — requirement-set document ──────────────────────────────────────
META_WIDTHS = [2700, 6660]  # dxa, sum 9360 == 6.5 in
REQ_WIDTHS = [1500, 7860]  # dxa, sum 9360 == 6.5 in


def _portrait(doc) -> None:
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.left_margin = sec.right_margin = Inches(1.0)
    sec.top_margin = sec.bottom_margin = Inches(1.0)


def _meta_table(doc, pairs) -> None:
    if not pairs:
        return
    table = _new_table(doc, rows=len(pairs), cols=2)
    _single_borders(table)
    for i, kv in enumerate(pairs):
        label_cell, value_cell = table.rows[i].cells
        _shade_cell(label_cell, FILL_META_LABEL)
        _style_cell_text(label_cell, kv.label, bold=True, size=9.0, name=None)
        _style_cell_text(value_cell, kv.value, bold=False, size=9.0, name=None)
    _fix_widths(table, META_WIDTHS)


def _req_table(doc, rows) -> None:
    table = _new_table(doc, rows=1, cols=2)
    _single_borders(table)
    for cell, text in zip(table.rows[0].cells, ("Req ID", "Requirement")):
        _shade_cell(cell, FILL_HEADER_B)
        _style_cell_text(cell, text, bold=True, size=9.0, color=WHITE, name=None)
    for row in rows:
        id_cell, text_cell = table.add_row().cells
        _style_cell_text(id_cell, row.req_id, bold=True, size=9.0, name=None)
        _style_cell_text(text_cell, row.text, bold=False, size=9.0, name=None)
    _fix_widths(table, REQ_WIDTHS)


def render_reqset_docx(doc_model: GeneratedDoc) -> bytes:
    """Render a Product/Hardware/Software/Labeling requirement doc (.docx)."""
    doc = Document()
    _portrait(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(doc_model.title)
    tr.font.bold = True
    tr.font.size = Pt(24)
    tr.font.color.rgb = NAVY_REQSET

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(doc_model.product_name)
    sr.font.bold = True
    sr.font.size = Pt(18)
    sr.font.color.rgb = GREY_SUBTITLE

    _meta_table(doc, doc_model.context_table)

    lead = "Purpose" if doc_model.doc_type == "product" else "Scope"
    body = doc_model.purpose or doc_model.scope or ""
    if body:
        doc.add_paragraph(f"{lead}: {body}")

    n = 0
    for section in doc_model.sections:
        if section.level == 2:
            doc.add_heading(section.title, level=2)
        else:
            n += 1
            doc.add_heading(f"{n}. {section.title}", level=1)
            # The product overview renders as a metadata table, not a req table.
            # If the model put requirement rows in the first section anyway,
            # fall through and render them too — rows must never be dropped.
            if doc_model.doc_type == "product" and doc_model.overview_table and n == 1:
                _meta_table(doc, doc_model.overview_table)
                if not section.rows:
                    continue
        if section.rows:
            _req_table(doc, section.rows)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
