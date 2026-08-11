"""Turn an uploaded scope document into plain text for the pipeline.

Supports the formats the upload control already advertises (.docx, .pdf, .md,
.txt). Tables matter here — the Tata Elxsi sample requirement docs carry almost
all of their content in tables, so a paragraph-only extraction would return an
almost-empty document.
"""
from __future__ import annotations

import io

MAX_BYTES = 10 * 1024 * 1024  # 10 MB — generous for a requirements doc.
MAX_CHARS = 120_000  # Trim runaway documents before they reach the model.


class ParseError(Exception):
    """Raised when a document can't be read or has no usable text."""


def _parse_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise ParseError("python-docx is not installed on the backend.") from exc

    try:
        doc = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ParseError(f"Could not read the .docx file: {exc}") from exc

    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]

    # Tables carry the requirement IDs and statements in these documents.
    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts)


def _parse_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise ParseError("pypdf is not installed on the backend.") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as exc:
        raise ParseError(f"Could not read the .pdf file: {exc}") from exc


def _parse_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ParseError("Could not decode the file as text.")


def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from an uploaded document.

    Raises ParseError for unsupported types, unreadable files, oversized files,
    or documents that yield no usable text.
    """
    if not data:
        raise ParseError("The uploaded file is empty.")
    if len(data) > MAX_BYTES:
        raise ParseError(
            f"File is {len(data) // (1024 * 1024)} MB; the limit is {MAX_BYTES // (1024 * 1024)} MB."
        )

    name = (filename or "").lower()
    if name.endswith(".docx"):
        text = _parse_docx(data)
    elif name.endswith(".pdf"):
        text = _parse_pdf(data)
    elif name.endswith((".md", ".txt", ".markdown")):
        text = _parse_text(data)
    elif name.endswith(".doc"):
        # Legacy binary .doc is a different format that python-docx cannot read.
        raise ParseError(
            "Legacy .doc files aren't supported — re-save as .docx and upload again."
        )
    else:
        raise ParseError(f"Unsupported file type: {filename}. Use .docx, .pdf, .md, or .txt.")

    text = text.strip()
    if len(text) < 50:
        raise ParseError(
            "No readable text found in the document — it may be a scan or image-only PDF."
        )

    return text[:MAX_CHARS]
